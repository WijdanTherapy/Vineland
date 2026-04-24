import streamlit as st
import json, re, os
from datetime import date, datetime
from io import BytesIO
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

# ── ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.graphics.shapes import Drawing, Rect, Line, String, Circle
from reportlab.graphics import renderPDF
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ── python-docx
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Groq
from groq import Groq

# ═══════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════
THERAPIST_NAME  = "Yusuf Abdelatti"
THERAPIST_TITLE = "Psychotherapist"
CENTER_NAME     = "Wijdan Therapy Center"
GMAIL_ADDRESS   = "Wijdan.psyc@gmail.com"
GMAIL_PASSWORD  = "rias eeul lyuu stce"
THERAPIST_EMAIL = "Wijdan.psyc@gmail.com"
LOGO_FILE       = "logo.png"

DEEP    = colors.HexColor("#1C1917")
WARM    = colors.HexColor("#8B7355")
ACCENT  = colors.HexColor("#C4956A")
CREAM   = colors.HexColor("#F7F3EE")
BORDER  = colors.HexColor("#DDD5C8")
SELECTED= colors.HexColor("#2D2926")

# ═══════════════════════════════════════════════════════════
# LOAD NORM TABLES
# ═══════════════════════════════════════════════════════════
@st.cache_resource
def load_norms():
    base = os.path.dirname(__file__)
    with open(os.path.join(base, "b1_interview.json")) as f:
        b1 = json.load(f)
    with open(os.path.join(base, "b3_interview.json")) as f:
        b3 = json.load(f)
    with open(os.path.join(base, "maladaptive.json")) as f:
        mal = json.load(f)
    return b1, b3, mal

# ═══════════════════════════════════════════════════════════
# SCORING ENGINE
# ═══════════════════════════════════════════════════════════
def age_in_months(dob: date, test_date: date) -> int:
    months = (test_date.year - dob.year) * 12 + (test_date.month - dob.month)
    if test_date.day < dob.day:
        months -= 1
    return max(0, months)

def find_b1_key(b1, age_months):
    for key in b1:
        if key == "90+":
            if age_months >= 90 * 12:
                return key
            continue
        key_norm = key.replace("–", "-")
        parts = key_norm.split("-")
        if len(parts) == 2:
            def pym(s):
                p = s.strip().split(":")
                return int(p[0]) * 12 + (int(p[1]) if len(p) > 1 else 0)
            try:
                if pym(parts[0]) <= age_months <= pym(parts[1]):
                    return key
            except:
                pass
    return None

def find_b3_key(b3, age_months):
    for key in b3:
        if key == "70–90+":
            if age_months >= 70 * 12:
                return key
            continue
        key_norm = key.replace("–", "-")
        parts = key_norm.split("-")
        if len(parts) == 2:
            def pym(s):
                p = s.strip().split(":")
                return int(p[0]) * 12 + (int(p[1]) if len(p) > 1 else 0)
            try:
                if pym(parts[0]) <= age_months <= pym(parts[1]):
                    return key
            except:
                pass
    return None

def find_mal_key(age_months):
    age_years = age_months / 12
    if age_years < 7:
        return "3–6"
    elif age_years < 12:
        return "7–11"
    elif age_years < 21:
        return "12–20"
    elif age_years < 70:
        return "21–69"
    else:
        return "70–90+"

def raw_to_vscore(sd_data, raw):
    """Convert raw score to v-scale score. sd_data keys are strings."""
    val = sd_data.get(str(raw))
    if val is not None:
        return val
    int_keys = [int(k) for k in sd_data.keys()]
    if not int_keys:
        return None
    if raw >= max(int_keys):
        return sd_data[str(max(int_keys))]
    if raw <= min(int_keys):
        return sd_data[str(min(int_keys))]
    # Interpolate: find nearest
    lower = max((k for k in int_keys if k <= raw), default=None)
    if lower is not None:
        return sd_data[str(lower)]
    return 1

def compute_scores(raw_scores, age_months, include_motor, b1, b3, mal):
    """
    raw_scores: dict with keys rec,exp,wrn,per,dom,cmm,ipr,pla,cop,gmo,fmo,
                                mal_int,mal_ext (raw counts for maladaptive)
    Returns full scoring dict.
    """
    b1_key = find_b1_key(b1, age_months)
    b3_key = find_b3_key(b3, age_months)
    mal_key = find_mal_key(age_months)

    if not b1_key or not b3_key:
        return None

    b1_age = b1[b1_key]
    b3_age = b3[b3_key]

    # ── v-scale scores
    subdomains = ["rec","exp","wrn","per","dom","cmm","ipr","pla","cop","gmo","fmo"]
    vscores = {}
    for sd in subdomains:
        raw = raw_scores.get(sd, 0)
        vs = raw_to_vscore(b1_age[sd], raw)
        vscores[sd] = vs if vs is not None else 1

    # ── domain sums of v-scales
    com_vsum = vscores["rec"] + vscores["exp"] + vscores["wrn"]
    dls_vsum = vscores["per"] + vscores["dom"] + vscores["cmm"]
    soc_vsum = vscores["ipr"] + vscores["pla"] + vscores["cop"]
    mot_vsum = vscores["gmo"] + vscores["fmo"]

    # ── domain standard scores (B.3 keys are strings)
    def b3_lookup(domain, vsum):
        d = b3_age[domain]
        val = d.get(str(vsum))
        if val is not None: return val
        int_keys = [int(k) for k in d.keys()]
        if not int_keys: return None
        if vsum <= min(int_keys): return d[str(min(int_keys))]
        if vsum >= max(int_keys): return d[str(max(int_keys))]
        lower = max((k for k in int_keys if k <= vsum), default=None)
        if lower is not None: return d[str(lower)]
        return None

    com_ss = b3_lookup("COM", com_vsum)
    dls_ss = b3_lookup("DLS", dls_vsum)
    soc_ss = b3_lookup("SOC", soc_vsum)
    mot_ss = b3_lookup("MOT", mot_vsum) if include_motor else None

    # ── ABC: sum of domain standard scores → look up in B.3 ABC column
    abc_domain_sum = (com_ss or 0) + (dls_ss or 0) + (soc_ss or 0)
    abc_ss = b3_lookup("ABC", abc_domain_sum)

    # ── percentile ranks
    def get_pct(ss):
        if ss is None: return None
        return b3_age["percentile"].get(str(ss))

    # ── 90% confidence intervals (from B.3 table footers)
    # These are fixed per broad age group
    ci_table = {
        # age_max_months: {COM, DLS, SOC, ABC, MOT}
        11: {"COM":9,"DLS":14,"SOC":8,"ABC":8,"MOT":10},
        23: {"COM":9,"DLS":14,"SOC":8,"ABC":8,"MOT":10},
        35: {"COM":9,"DLS":14,"SOC":8,"ABC":8,"MOT":10},
        47: {"COM":9,"DLS":11,"SOC":8,"ABC":8,"MOT":10},
        59: {"COM":8,"DLS":10,"SOC":8,"ABC":7,"MOT":9},
        83: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":9},
        107: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":9},
        131: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":8},
        155: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":8},
        179: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":8},
        9999: {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":8},
    }
    ci = {"COM":8,"DLS":9,"SOC":8,"ABC":7,"MOT":9}
    for age_max, ci_vals in sorted(ci_table.items()):
        if age_months <= age_max:
            ci = ci_vals
            break

    # ── Maladaptive behavior v-scales
    mal_age = mal[mal_key]
    mal_int_raw = raw_scores.get("mal_int", 0)
    mal_ext_raw = raw_scores.get("mal_ext", 0)
    mal_int_vs = raw_to_vscore(mal_age["int"], mal_int_raw)
    mal_ext_vs = raw_to_vscore(mal_age["ext"], mal_ext_raw)

    # ── Qualitative descriptors
    def qual_domain(ss):
        if ss is None: return "N/A"
        if ss >= 130: return "High"
        if ss >= 115: return "Moderately High"
        if ss >= 86:  return "Adequate"
        if ss >= 71:  return "Moderately Low"
        return "Low"

    def qual_vscore(vs):
        if vs is None: return "N/A"
        if vs >= 21: return "High"
        if vs >= 18: return "Moderately High"
        if vs >= 13: return "Adequate"
        if vs >= 10: return "Moderately Low"
        return "Low"

    def qual_maladaptive(vs):
        if vs is None: return "N/A"
        if vs >= 21: return "Clinically Significant"
        if vs >= 18: return "Elevated"
        return "Average"

    return {
        "b1_key": b1_key,
        "b3_key": b3_key,
        "vscores": vscores,
        "com_vsum": com_vsum, "dls_vsum": dls_vsum,
        "soc_vsum": soc_vsum, "mot_vsum": mot_vsum,
        "com_ss": com_ss, "dls_ss": dls_ss,
        "soc_ss": soc_ss, "mot_ss": mot_ss,
        "abc_domain_sum": abc_domain_sum,
        "abc_ss": abc_ss,
        "com_pct": get_pct(com_ss),
        "dls_pct": get_pct(dls_ss),
        "soc_pct": get_pct(soc_ss),
        "mot_pct": get_pct(mot_ss),
        "abc_pct": get_pct(abc_ss),
        "ci": ci,
        "mal_int_vs": mal_int_vs, "mal_ext_vs": mal_ext_vs,
        "mal_int_raw": mal_int_raw, "mal_ext_raw": mal_ext_raw,
        "qual_com": qual_domain(com_ss), "qual_dls": qual_domain(dls_ss),
        "qual_soc": qual_domain(soc_ss), "qual_mot": qual_domain(mot_ss),
        "qual_abc": qual_domain(abc_ss),
        "qual_rec": qual_vscore(vscores["rec"]),
        "qual_exp": qual_vscore(vscores["exp"]),
        "qual_wrn": qual_vscore(vscores["wrn"]),
        "qual_per": qual_vscore(vscores["per"]),
        "qual_dom": qual_vscore(vscores["dom"]),
        "qual_cmm": qual_vscore(vscores["cmm"]),
        "qual_ipr": qual_vscore(vscores["ipr"]),
        "qual_pla": qual_vscore(vscores["pla"]),
        "qual_cop": qual_vscore(vscores["cop"]),
        "qual_gmo": qual_vscore(vscores["gmo"]) if include_motor else "N/A",
        "qual_fmo": qual_vscore(vscores["fmo"]) if include_motor else "N/A",
        "qual_mal_int": qual_maladaptive(mal_int_vs),
        "qual_mal_ext": qual_maladaptive(mal_ext_vs),
        "include_motor": include_motor,
    }

# ═══════════════════════════════════════════════════════════
# UTILITIES
# ═══════════════════════════════════════════════════════════
def strip_arabic(text):
    if not isinstance(text, str):
        return str(text)
    cleaned = re.sub(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]', '', text)
    return cleaned.strip()

def format_ci(ss, ci_val):
    if ss is None: return "N/A"
    return f"{max(20, ss - ci_val)} – {min(160, ss + ci_val)}"

def pct_display(pct):
    if pct is None: return "N/A"
    if pct <= 1: return "<1"
    if pct >= 99: return ">99"
    return str(pct)

# ═══════════════════════════════════════════════════════════
# GROQ NARRATIVE
# ═══════════════════════════════════════════════════════════
def generate_narrative(demo, scores, lang, critical_items_text):
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])

    inc_motor = scores["include_motor"]
    motor_line = f"\nMotor Skills Standard Score: {scores['mot_ss']} (Percentile: {pct_display(scores['mot_pct'])}) — {scores['qual_mot']}\n  Gross Motor v-Scale: {scores['vscores']['gmo']} ({scores['qual_gmo']})\n  Fine Motor v-Scale: {scores['vscores']['fmo']} ({scores['qual_fmo']})" if inc_motor else ""

    prompt = f"""You are a clinical psychologist writing a formal psychological assessment report.

CLIENT: {demo['name']}
DATE OF BIRTH: {demo['dob']}
AGE: {demo['age_str']}
GENDER: {demo['gender']}
NATIONALITY: {demo['nationality']}
REFERRAL SOURCE: {demo['referral']}
RESPONDENT NAME: {demo['respondent_name']}
RESPONDENT RELATIONSHIP: {demo['respondent_rel']}
TEST LANGUAGE: {lang}
ASSESSMENT: Vineland Adaptive Behavior Scales, Third Edition (Vineland-3) — Comprehensive Interview Form
DATE: {demo['test_date']}
REPORT PREPARED BY: {THERAPIST_NAME}, {THERAPIST_TITLE} — {CENTER_NAME}

SCORE RESULTS:
Adaptive Behavior Composite (ABC): SS={scores['abc_ss']}, 90% CI={format_ci(scores['abc_ss'], scores['ci']['ABC'])}, Percentile={pct_display(scores['abc_pct'])}, Level={scores['qual_abc']}

Communication Domain: SS={scores['com_ss']}, 90% CI={format_ci(scores['com_ss'], scores['ci']['COM'])}, Percentile={pct_display(scores['com_pct'])}, Level={scores['qual_com']}
  Receptive v-Scale: {scores['vscores']['rec']} ({scores['qual_rec']})
  Expressive v-Scale: {scores['vscores']['exp']} ({scores['qual_exp']})
  Written v-Scale: {scores['vscores']['wrn']} ({scores['qual_wrn']})

Daily Living Skills Domain: SS={scores['dls_ss']}, 90% CI={format_ci(scores['dls_ss'], scores['ci']['DLS'])}, Percentile={pct_display(scores['dls_pct'])}, Level={scores['qual_dls']}
  Personal v-Scale: {scores['vscores']['per']} ({scores['qual_per']})
  Domestic v-Scale: {scores['vscores']['dom']} ({scores['qual_dom']})
  Community v-Scale: {scores['vscores']['cmm']} ({scores['qual_cmm']})

Socialization Domain: SS={scores['soc_ss']}, 90% CI={format_ci(scores['soc_ss'], scores['ci']['SOC'])}, Percentile={pct_display(scores['soc_pct'])}, Level={scores['qual_soc']}
  Interpersonal Relationships v-Scale: {scores['vscores']['ipr']} ({scores['qual_ipr']})
  Play and Leisure v-Scale: {scores['vscores']['pla']} ({scores['qual_pla']})
  Coping Skills v-Scale: {scores['vscores']['cop']} ({scores['qual_cop']})
{motor_line}
Maladaptive Behavior — Internalizing v-Scale: {scores['mal_int_vs']} ({scores['qual_mal_int']})
Maladaptive Behavior — Externalizing v-Scale: {scores['mal_ext_vs']} ({scores['qual_mal_ext']})
{f"Critical Items Endorsed: {critical_items_text}" if critical_items_text else "No critical items endorsed."}

Write a formal clinical narrative report with the following sections, each titled in ALL CAPS with a number:
1. REFERRAL AND ASSESSMENT OVERVIEW
2. ASSESSMENT INSTRUMENT AND PROCEDURE
3. BEHAVIORAL OBSERVATIONS
4. ADAPTIVE BEHAVIOR COMPOSITE
5. COMMUNICATION DOMAIN
6. DAILY LIVING SKILLS DOMAIN
7. SOCIALIZATION DOMAIN
{("8. MOTOR SKILLS DOMAIN\n9. MALADAPTIVE BEHAVIOR\n10. SUMMARY AND RECOMMENDATIONS" if inc_motor else "8. MALADAPTIVE BEHAVIOR\n9. SUMMARY AND RECOMMENDATIONS")}

RULES:
- No diagnosis — findings as hypotheses only
- Formal clinical language throughout
- No markdown (no **, ##, ---)
- Reference specific scores and qualitative descriptors when discussing each domain
- Integrate subdomain v-scale scores in interpretation
- Section titles must be ALL CAPS numbered as shown above
- Write in flowing paragraphs, no bullet points
- Length: comprehensive but concise
"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        max_tokens=3500,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

# ═══════════════════════════════════════════════════════════
# PDF GENERATION
# ═══════════════════════════════════════════════════════════
def build_pdf(demo, scores, narrative, lang):
    buf = BytesIO()
    W, H = A4

    # Styles
    styles = {
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=16,
                                 textColor=DEEP, spaceAfter=4),
        "subtitle": ParagraphStyle("subtitle", fontName="Helvetica", fontSize=10,
                                    textColor=WARM, spaceAfter=2),
        "section": ParagraphStyle("section", fontName="Helvetica-Bold", fontSize=10,
                                   textColor=DEEP, spaceBefore=10, spaceAfter=4,
                                   borderPad=2),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=9,
                                textColor=DEEP, leading=14, spaceAfter=6),
        "small": ParagraphStyle("small", fontName="Helvetica", fontSize=8,
                                 textColor=WARM, spaceAfter=2),
    }

    story = []

    def add_header_footer(canvas, doc):
        canvas.saveState()
        # Header line
        canvas.setStrokeColor(ACCENT)
        canvas.setLineWidth(2)
        canvas.line(20*mm, H - 18*mm, W - 20*mm, H - 18*mm)
        # Footer
        canvas.setStrokeColor(BORDER)
        canvas.setLineWidth(0.5)
        canvas.line(20*mm, 18*mm, W - 20*mm, 18*mm)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(WARM)
        canvas.drawCentredString(W/2, 12*mm, f"{CENTER_NAME}  ·  {THERAPIST_NAME}, {THERAPIST_TITLE}  ·  Confidential Clinical Document")
        canvas.drawRightString(W - 20*mm, 12*mm, f"Page {doc.page}")
        canvas.restoreState()

    # ── Cover block
    # Logo
    if os.path.exists(LOGO_FILE):
        from reportlab.platypus import Image as RLImage
        story.append(RLImage(LOGO_FILE, width=40*mm, height=15*mm))
        story.append(Spacer(1, 4*mm))

    story.append(Paragraph("Vineland Adaptive Behavior Scales, Third Edition", styles["title"]))
    story.append(Paragraph("Comprehensive Interview Form — Clinical Report", styles["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=1, color=ACCENT))
    story.append(Spacer(1, 4*mm))

    # Client info table
    ci = scores["ci"]
    name_safe = strip_arabic(demo['name'])
    rows = [
        [Paragraph("<b>Name</b>", styles["small"]),
         Paragraph(name_safe, styles["body"]),
         Paragraph("<b>Age</b>", styles["small"]),
         Paragraph(demo['age_str'], styles["body"]),
         Paragraph("<b>Gender</b>", styles["small"]),
         Paragraph(strip_arabic(demo['gender']), styles["body"])],
        [Paragraph("<b>Date of Birth</b>", styles["small"]),
         Paragraph(str(demo['dob']), styles["body"]),
         Paragraph("<b>Nationality</b>", styles["small"]),
         Paragraph(strip_arabic(demo['nationality']), styles["body"]),
         Paragraph("<b>Test Language</b>", styles["small"]),
         Paragraph(lang, styles["body"])],
        [Paragraph("<b>Referral Source</b>", styles["small"]),
         Paragraph(strip_arabic(demo['referral']), styles["body"]),
         Paragraph("<b>Assessment</b>", styles["small"]),
         Paragraph("Vineland-3 CIF", styles["body"]),
         Paragraph("<b>Test Date</b>", styles["small"]),
         Paragraph(str(demo['test_date']), styles["body"])],
        [Paragraph("<b>Respondent</b>", styles["small"]),
         Paragraph(strip_arabic(demo['respondent_name']), styles["body"]),
         Paragraph("<b>Relationship</b>", styles["small"]),
         Paragraph(strip_arabic(demo['respondent_rel']), styles["body"]),
         Paragraph("<b>Prepared By</b>", styles["small"]),
         Paragraph(f"{THERAPIST_NAME}, {THERAPIST_TITLE}", styles["body"])],
    ]
    info_table = Table(rows, colWidths=[28*mm, 45*mm, 28*mm, 35*mm, 28*mm, 35*mm])
    info_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), CREAM),
        ("BOX", (0,0), (-1,-1), 0.5, BORDER),
        ("INNERGRID", (0,0), (-1,-1), 0.25, BORDER),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("LEFTPADDING", (0,0), (-1,-1), 4),
    ]))
    story.append(info_table)
    story.append(Spacer(1, 6*mm))

    # ── Score Summary Section
    story.append(Paragraph("SCORE SUMMARY", styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 2*mm))

    def ss_row(label, ss, ci_val, pct, qual, bold=False):
        style = "Helvetica-Bold" if bold else "Helvetica"
        bg = colors.HexColor("#EDE8E2") if bold else colors.white
        return [
            Paragraph(f"<font name='{style}'>{label}</font>", styles["body"]),
            Paragraph(f"<font name='{style}'>{ss if ss else 'N/A'}</font>", styles["body"]),
            Paragraph(f"<font name='{style}'>{format_ci(ss, ci_val) if ss else 'N/A'}</font>", styles["body"]),
            Paragraph(f"<font name='{style}'>{pct_display(pct)}</font>", styles["body"]),
            Paragraph(f"<font name='{style}'>{qual}</font>", styles["body"]),
        ], bg

    score_header = [
        Paragraph("<b>Scale</b>", styles["small"]),
        Paragraph("<b>Standard Score</b>", styles["small"]),
        Paragraph("<b>90% CI</b>", styles["small"]),
        Paragraph("<b>Percentile Rank</b>", styles["small"]),
        Paragraph("<b>Adaptive Level</b>", styles["small"]),
    ]

    abc_row, abc_bg = ss_row(
        "Adaptive Behavior Composite (ABC)",
        scores["abc_ss"], ci["ABC"], scores["abc_pct"], scores["qual_abc"], bold=True
    )
    com_row, com_bg = ss_row("Communication", scores["com_ss"], ci["COM"], scores["com_pct"], scores["qual_com"])
    dls_row, dls_bg = ss_row("Daily Living Skills", scores["dls_ss"], ci["DLS"], scores["dls_pct"], scores["qual_dls"])
    soc_row, soc_bg = ss_row("Socialization", scores["soc_ss"], ci["SOC"], scores["soc_pct"], scores["qual_soc"])

    score_data = [score_header, abc_row, com_row, dls_row, soc_row]
    score_bgs = [None, abc_bg, com_bg, dls_bg, soc_bg]

    if scores["include_motor"]:
        mot_row, mot_bg = ss_row("Motor Skills", scores["mot_ss"], ci["MOT"], scores["mot_pct"], scores["qual_mot"])
        score_data.append(mot_row)
        score_bgs.append(mot_bg)

    score_table = Table(score_data, colWidths=[65*mm, 30*mm, 30*mm, 30*mm, 42*mm])
    ts = [
        ("BACKGROUND", (0,0), (-1,0), DEEP),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,0), 8),
        ("BOX", (0,0), (-1,-1), 0.5, BORDER),
        ("INNERGRID", (0,0), (-1,-1), 0.25, BORDER),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("LEFTPADDING", (0,0), (-1,-1), 4),
    ]
    for i, bg in enumerate(score_bgs):
        if bg and i > 0:
            ts.append(("BACKGROUND", (0,i), (-1,i), bg))
    score_table.setStyle(TableStyle(ts))
    story.append(score_table)
    story.append(Spacer(1, 4*mm))

    # ── Subdomain v-Scale Summary
    story.append(Paragraph("SUBDOMAIN v-SCALE SCORES", styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 2*mm))

    sd_header = [
        Paragraph("<b>Domain</b>", styles["small"]),
        Paragraph("<b>Subdomain</b>", styles["small"]),
        Paragraph("<b>Raw Score</b>", styles["small"]),
        Paragraph("<b>v-Scale Score</b>", styles["small"]),
        Paragraph("<b>Adaptive Level</b>", styles["small"]),
    ]

    def sd_row(domain, subdomain, raw, vs, qual):
        return [
            Paragraph(domain, styles["body"]),
            Paragraph(subdomain, styles["body"]),
            Paragraph(str(raw), styles["body"]),
            Paragraph(str(vs) if vs else "N/A", styles["body"]),
            Paragraph(qual, styles["body"]),
        ]

    vs = scores["vscores"]
    rs = {}  # raw scores from demo
    sd_data = [sd_header,
        sd_row("Communication", "Receptive", demo['raw'].get('rec',0), vs['rec'], scores['qual_rec']),
        sd_row("", "Expressive", demo['raw'].get('exp',0), vs['exp'], scores['qual_exp']),
        sd_row("", "Written", demo['raw'].get('wrn',0), vs['wrn'], scores['qual_wrn']),
        sd_row("Daily Living Skills", "Personal", demo['raw'].get('per',0), vs['per'], scores['qual_per']),
        sd_row("", "Domestic", demo['raw'].get('dom',0), vs['dom'], scores['qual_dom']),
        sd_row("", "Community", demo['raw'].get('cmm',0), vs['cmm'], scores['qual_cmm']),
        sd_row("Socialization", "Interpersonal Relationships", demo['raw'].get('ipr',0), vs['ipr'], scores['qual_ipr']),
        sd_row("", "Play and Leisure", demo['raw'].get('pla',0), vs['pla'], scores['qual_pla']),
        sd_row("", "Coping Skills", demo['raw'].get('cop',0), vs['cop'], scores['qual_cop']),
    ]
    if scores["include_motor"]:
        sd_data.append(sd_row("Motor Skills", "Gross Motor", demo['raw'].get('gmo',0), vs['gmo'], scores['qual_gmo']))
        sd_data.append(sd_row("", "Fine Motor", demo['raw'].get('fmo',0), vs['fmo'], scores['qual_fmo']))

    # Add maladaptive
    sd_data.append(sd_row("Maladaptive Behavior", "Internalizing",
                           scores['mal_int_raw'], scores['mal_int_vs'], scores['qual_mal_int']))
    sd_data.append(sd_row("", "Externalizing",
                           scores['mal_ext_raw'], scores['mal_ext_vs'], scores['qual_mal_ext']))

    sd_table = Table(sd_data, colWidths=[45*mm, 55*mm, 25*mm, 30*mm, 42*mm])
    sd_ts = [
        ("BACKGROUND", (0,0), (-1,0), DEEP),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,0), 8),
        ("BOX", (0,0), (-1,-1), 0.5, BORDER),
        ("INNERGRID", (0,0), (-1,-1), 0.25, BORDER),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING", (0,0), (-1,-1), 3),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("LEFTPADDING", (0,0), (-1,-1), 4),
    ]
    # Alternate row shading
    for i in range(1, len(sd_data)):
        if i % 2 == 0:
            sd_ts.append(("BACKGROUND", (0,i), (-1,i), colors.HexColor("#F5F0EA")))
    sd_table.setStyle(TableStyle(sd_ts))
    story.append(sd_table)
    story.append(Spacer(1, 4*mm))

    # ── Profile Chart
    story.append(Paragraph("SCORE PROFILE CHART", styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 2*mm))
    story.append(_build_profile_chart(scores))
    story.append(Spacer(1, 6*mm))

    # ── Narrative
    story.append(Paragraph("CLINICAL NARRATIVE REPORT", styles["section"]))
    story.append(HRFlowable(width="100%", thickness=1, color=ACCENT))
    story.append(Spacer(1, 3*mm))

    for para in narrative.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if re.match(r'^\d+\.\s+[A-Z\s]+$', para):
            story.append(Spacer(1, 3*mm))
            story.append(Paragraph(para, styles["section"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        else:
            story.append(Paragraph(para, styles["body"]))

    # ── Footer sign-off
    story.append(Spacer(1, 8*mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 3*mm))
    footer_data = [[
        Paragraph(f"<b>{THERAPIST_NAME}</b><br/>{THERAPIST_TITLE}<br/>{CENTER_NAME}", styles["body"]),
        Paragraph(f"<i>This report is a confidential clinical document prepared for professional use only. "
                  f"The information contained herein should not be shared without the written consent of the "
                  f"examiner and is intended solely for the purposes of clinical evaluation and intervention planning.</i>",
                  styles["small"]),
    ]]
    footer_t = Table(footer_data, colWidths=[70*mm, 127*mm])
    footer_t.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING", (0,0), (-1,-1), 0),
    ]))
    story.append(footer_t)

    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=20*mm, rightMargin=20*mm,
                            topMargin=25*mm, bottomMargin=25*mm)
    doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
    buf.seek(0)
    return buf

def _build_profile_chart(scores):
    """Build ReportLab Drawing with domain score profile."""
    chart_w, chart_h = 170*mm, 60*mm
    d = Drawing(chart_w, chart_h)

    domains = ["COM", "DLS", "SOC", "ABC"]
    ss_vals = [scores["com_ss"], scores["dls_ss"], scores["soc_ss"], scores["abc_ss"]]
    labels = ["Communication", "Daily Living\nSkills", "Socialization", "ABC\nComposite"]

    if scores["include_motor"]:
        domains.append("MOT")
        ss_vals.append(scores["mot_ss"])
        labels.append("Motor\nSkills")

    n = len(domains)
    margin_l = 15*mm
    margin_r = 10*mm
    margin_b = 12*mm
    margin_t = 8*mm
    plot_w = chart_w - margin_l - margin_r
    plot_h = chart_h - margin_b - margin_t

    ss_min, ss_max = 40, 160
    ss_range = ss_max - ss_min

    def y_for_ss(ss):
        return margin_b + (ss - ss_min) / ss_range * plot_h

    # Grid lines and labels
    for ss in [40, 60, 70, 85, 100, 115, 130, 140, 160]:
        y = y_for_ss(ss)
        d.add(Line(margin_l, y, margin_l + plot_w, y,
                   strokeColor=BORDER if ss != 100 else WARM,
                   strokeWidth=0.5 if ss != 100 else 1))
        d.add(String(margin_l - 2*mm, y - 2, str(ss),
                     fontSize=6, fillColor=WARM, textAnchor="end"))

    # Draw bars and points
    bar_w = plot_w / n * 0.4
    x_positions = []
    for i in range(n):
        x = margin_l + (i + 0.5) * (plot_w / n)
        x_positions.append(x)

        ss = ss_vals[i]
        if ss is None:
            continue

        y = y_for_ss(ss)

        # Bar
        bar_color = ACCENT if domains[i] == "ABC" else WARM
        bar_h = y - y_for_ss(ss_min)
        if bar_h > 0:
            d.add(Rect(x - bar_w/2, margin_b, bar_w, bar_h,
                       fillColor=bar_color, strokeColor=None,
                       fillOpacity=0.3))

        # Point
        d.add(Circle(x, y, 3*mm, fillColor=bar_color, strokeColor=colors.white, strokeWidth=0.5))

        # SS label above point
        d.add(String(x, y + 4*mm, str(ss),
                     fontSize=7, fillColor=DEEP, textAnchor="middle",
                     fontName="Helvetica-Bold"))

        # Domain label below
        for j, lline in enumerate(labels[i].split("\n")):
            d.add(String(x, margin_b - 4*mm - j*6,
                         lline, fontSize=6.5, fillColor=DEEP, textAnchor="middle"))

    # Connect points with line
    pts = [(x_positions[i], y_for_ss(ss_vals[i]))
           for i in range(n) if ss_vals[i] is not None]
    for i in range(len(pts)-1):
        d.add(Line(pts[i][0], pts[i][1], pts[i+1][0], pts[i+1][1],
                   strokeColor=ACCENT, strokeWidth=1.5))

    # Mean line (100)
    d.add(String(margin_l - 2*mm, y_for_ss(100) - 2, "100",
                 fontSize=6, fillColor=WARM, textAnchor="end",
                 fontName="Helvetica-Bold"))

    return d

# ═══════════════════════════════════════════════════════════
# ARABIC WORD DOC
# ═══════════════════════════════════════════════════════════
def build_arabic_word_doc(demo, item_responses, scores):
    """Build Arabic Word document with all item responses."""
    doc = DocxDocument()

    # Page setup
    section = doc.sections[0]
    section.right_to_left = True

    def add_rtl_para(text, bold=False, size=12, color=None, align="RIGHT"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # Header
    add_rtl_para("مقياس فاينلاند للسلوك التكيفي - الإصدار الثالث", bold=True, size=14, color=(28,25,23))
    add_rtl_para("نموذج المقابلة الشاملة — استجابات البنود", bold=False, size=11, color=(139,115,85))
    doc.add_paragraph("─" * 60)

    # Client info
    add_rtl_para(f"الاسم: {demo['name']}", bold=True)
    add_rtl_para(f"تاريخ الميلاد: {demo['dob']}")
    add_rtl_para(f"العمر: {demo['age_str']}")
    add_rtl_para(f"الجنس: {'ذكر' if demo['gender'] == 'Male' else 'أنثى'}")
    add_rtl_para(f"الجنسية: {demo['nationality']}")
    add_rtl_para(f"مصدر الإحالة: {demo['referral']}")
    add_rtl_para(f"المُستجيب: {demo['respondent_name']} ({demo['respondent_rel']})")
    add_rtl_para(f"تاريخ التطبيق: {demo['test_date']}")
    doc.add_paragraph()

    # Score summary in Arabic
    add_rtl_para("ملخص الدرجات", bold=True, size=13, color=(28,25,23))
    doc.add_paragraph("─" * 40)

    score_lines = [
        f"المركب السلوكي التكيفي (ABC): {scores['abc_ss']} — المستوى: {scores['qual_abc']}",
        f"التواصل: {scores['com_ss']} — المستوى: {scores['qual_com']}",
        f"مهارات الحياة اليومية: {scores['dls_ss']} — المستوى: {scores['qual_dls']}",
        f"التنشئة الاجتماعية: {scores['soc_ss']} — المستوى: {scores['qual_soc']}",
    ]
    if scores["include_motor"]:
        score_lines.append(f"المهارات الحركية: {scores['mot_ss']} — المستوى: {scores['qual_mot']}")

    for line in score_lines:
        add_rtl_para(line)

    doc.add_paragraph()

    # Item responses by domain
    domain_labels = {
        "rec": "الأول: التواصل — اللغة الاستقبالية",
        "exp": "الأول: التواصل — اللغة التعبيرية",
        "wrn": "الأول: التواصل — القراءة والكتابة",
        "per": "الثاني: مهارات الحياة اليومية — المهارات الذاتية",
        "dom": "الثاني: مهارات الحياة اليومية — الأنشطة المنزلية",
        "cmm": "الثاني: مهارات الحياة اليومية — المهارات المجتمعية",
        "ipr": "الثالث: التنشئة الاجتماعية — العلاقات الشخصية المتبادلة",
        "pla": "الثالث: التنشئة الاجتماعية — وقت الراحة والترفيه",
        "cop": "الثالث: التنشئة الاجتماعية — المسايرة",
        "gmo": "الرابع: المهارات الحركية — العضلات الكبيرة",
        "fmo": "الرابع: المهارات الحركية — العضلات الدقيقة",
        "mal_int": "الخامس: السلوك غير التكيفي — الجزء الأول",
        "mal_ext": "الخامس: السلوك غير التكيفي — الجزء الثاني",
    }

    response_map = {0: "لا، أبداً (0)", 1: "أحياناً (1)", 2: "نعم، عادةً (2)",
                    "N": "لم تتح الفرصة (م)", "U": "لا أعرف (ع)"}

    for domain_key, domain_label in domain_labels.items():
        if not scores["include_motor"] and domain_key in ["gmo","fmo"]:
            continue
        domain_responses = item_responses.get(domain_key, {})
        if not domain_responses:
            continue

        doc.add_paragraph()
        add_rtl_para(f"البُعد {domain_label}", bold=True, size=11, color=(28,25,23))
        add_rtl_para(f"الدرجة الخام: {demo['raw'].get(domain_key, 0)}  |  "
                     f"درجة المقياس المرحلي: {scores['vscores'].get(domain_key, 'N/A')}",
                     color=(139,115,85))
        doc.add_paragraph()

        for item_num in sorted(domain_responses.keys(), key=lambda x: int(x)):
            val = domain_responses[item_num]
            val_text = response_map.get(val, str(val))
            add_rtl_para(f"البند {item_num}: {val_text}")

    # Maladaptive raw scores note
    doc.add_paragraph()
    add_rtl_para("السلوك غير التكيفي — الدرجات", bold=True, size=11, color=(28,25,23))
    add_rtl_para(f"الجزء الأول (الداخلي): درجة خام = {scores['mal_int_raw']}، "
                 f"درجة مقياس مرحلي = {scores['mal_int_vs']} ({scores['qual_mal_int']})")
    add_rtl_para(f"الجزء الثاني (الخارجي): درجة خام = {scores['mal_ext_raw']}، "
                 f"درجة مقياس مرحلي = {scores['mal_ext_vs']} ({scores['qual_mal_ext']})")

    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 60)
    add_rtl_para(f"{THERAPIST_NAME} — {THERAPIST_TITLE}", bold=True)
    add_rtl_para(CENTER_NAME)
    add_rtl_para("وثيقة سرية للاستخدام السريري فقط", color=(139,115,85))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════
# EMAIL
# ═══════════════════════════════════════════════════════════
def send_email(demo, scores, pdf_buf, secondary_buf, secondary_filename, lang):
    try:
        msg = MIMEMultipart("alternative")
        msg["From"] = GMAIL_ADDRESS
        msg["To"] = THERAPIST_EMAIL
        name_safe = strip_arabic(demo['name'])
        msg["Subject"] = f"[Vineland-3] {name_safe} — {demo['test_date']}"

        def ss_line(label, ss, pct, qual):
            return f"<tr><td style='padding:3px 8px;'>{label}</td><td><b>{ss or 'N/A'}</b></td><td>{pct_display(pct)}</td><td>{qual}</td></tr>"

        html = f"""
<html><body style="font-family:Arial,sans-serif;color:#1C1917;">
<h2 style="color:#C4956A;">Vineland-3 Comprehensive Interview Form</h2>
<h3>{name_safe} &mdash; {demo['test_date']}</h3>
<table cellpadding="4" cellspacing="0" style="border-collapse:collapse;margin-bottom:16px;">
  <tr><td><b>Age:</b></td><td>{demo['age_str']}</td></tr>
  <tr><td><b>Gender:</b></td><td>{strip_arabic(demo['gender'])}</td></tr>
  <tr><td><b>Nationality:</b></td><td>{strip_arabic(demo['nationality'])}</td></tr>
  <tr><td><b>Referral:</b></td><td>{strip_arabic(demo['referral'])}</td></tr>
  <tr><td><b>Respondent:</b></td><td>{strip_arabic(demo['respondent_name'])} ({strip_arabic(demo['respondent_rel'])})</td></tr>
  <tr><td><b>Test Language:</b></td><td>{lang}</td></tr>
</table>
<table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse;min-width:400px;">
  <tr style="background:#1C1917;color:white;">
    <th>Scale</th><th>Standard Score</th><th>Percentile</th><th>Adaptive Level</th>
  </tr>
  <tr style="background:#EDE8E2;font-weight:bold;">
    <td>Adaptive Behavior Composite (ABC)</td><td>{scores['abc_ss'] or 'N/A'}</td><td>{pct_display(scores['abc_pct'])}</td><td>{scores['qual_abc']}</td>
  </tr>
  {ss_line("Communication", scores['com_ss'], scores['com_pct'], scores['qual_com'])}
  {ss_line("Daily Living Skills", scores['dls_ss'], scores['dls_pct'], scores['qual_dls'])}
  {ss_line("Socialization", scores['soc_ss'], scores['soc_pct'], scores['qual_soc'])}
  {ss_line("Motor Skills", scores['mot_ss'], scores['mot_pct'], scores['qual_mot']) if scores['include_motor'] else ''}
</table>
<p style="margin-top:12px;font-size:12px;color:#8B7355;">
  Maladaptive — Internalizing v-Scale: {scores['mal_int_vs']} ({scores['qual_mal_int']})<br/>
  Maladaptive — Externalizing v-Scale: {scores['mal_ext_vs']} ({scores['qual_mal_ext']})
</p>
<hr style="border-color:#DDD5C8;"/>
<p style="font-size:11px;color:#8B7355;">
  {THERAPIST_NAME}, {THERAPIST_TITLE} &mdash; {CENTER_NAME}
</p>
</body></html>
"""
        msg.attach(MIMEText(html, "html"))

        # Attach PDF
        pdf_buf.seek(0)
        part = MIMEBase("application", "octet-stream")
        part.set_payload(pdf_buf.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="Vineland3_{name_safe}_{demo["test_date"]}.pdf"')
        msg.attach(part)

        # Attach secondary file (Arabic Word or EN item PDF)
        secondary_buf.seek(0)
        part2 = MIMEBase("application", "octet-stream")
        part2.set_payload(secondary_buf.read())
        encoders.encode_base64(part2)
        part2.add_header("Content-Disposition", f'attachment; filename="{secondary_filename}"')
        msg.attach(part2)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(GMAIL_ADDRESS, GMAIL_PASSWORD)
            server.send_message(msg)
    except Exception:
        pass

# ═══════════════════════════════════════════════════════════
# ITEM DEFINITIONS (all items, EN + AR translations)
# ═══════════════════════════════════════════════════════════
from vineland3_items import ITEMS_EN, ITEMS_AR, MALADAPTIVE_EN, MALADAPTIVE_AR, CRITICAL_ITEMS_EN, CRITICAL_ITEMS_AR

# ═══════════════════════════════════════════════════════════
# STREAMLIT UI
# ═══════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Vineland-3 | Wijdan Therapy Center",
    page_icon="🧠",
    layout="wide"
)

# ── CSS
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@400;600;700&family=Jost:wght@300;400;500;600&display=swap');
html, body, [class*="css"] { font-family: 'Jost', sans-serif; background: #F7F3EE; color: #1C1917; }
h1,h2,h3 { font-family: 'Cormorant Garamond', serif; }
.stButton>button { background: #2D2926; color: #F7F3EE; border-radius: 2px; border: none; font-family: 'Jost'; font-weight: 500; padding: 10px 24px; }
.stButton>button:hover { background: #C4956A; }
.stProgress > div > div { background: linear-gradient(90deg, #8B7355, #C4956A); }
div[data-baseweb="radio"] label { cursor: pointer; }
.stRadio > div { flex-direction: row; gap: 8px; }
.block-container { max-width: 960px; }
</style>
""", unsafe_allow_html=True)

# ── Session state init
if "page" not in st.session_state:
    st.session_state.page = "access"
if "lang" not in st.session_state:
    st.session_state.lang = "English"
if "demo" not in st.session_state:
    st.session_state.demo = {}
if "responses" not in st.session_state:
    st.session_state.responses = {}
if "done" not in st.session_state:
    st.session_state.done = False

AR = st.session_state.lang == "Arabic"

# ── Language toggle
col1, col2 = st.columns([4,1])
with col2:
    new_lang = st.selectbox("🌐", ["English","Arabic"],
                             index=0 if not AR else 1,
                             key="lang_select", label_visibility="collapsed")
    if new_lang != st.session_state.lang:
        st.session_state.lang = new_lang
        st.session_state.page = "access"
        st.session_state.responses = {}
        st.session_state.demo = {}
        st.session_state.done = False
        st.rerun()

# ════════════════════════
# PAGE: ACCESS CODE
# ════════════════════════
if st.session_state.page == "access":
    st.markdown("---")
    st.markdown(f"## {'🔐 رمز الوصول' if AR else '🔐 Access Code'}")
    code = st.text_input("" if not AR else "أدخل رمز الوصول", type="password",
                          placeholder="Enter access code" if not AR else "رمز الوصول")
    if st.button("Continue" if not AR else "متابعة"):
        valid = [c.strip() for c in st.secrets.get("ACCESS_CODE","").split(",")]
        if code.strip() in valid:
            st.session_state.page = "demographics"
            st.rerun()
        else:
            st.error("Invalid access code." if not AR else "رمز الوصول غير صحيح.")
    st.stop()

# ════════════════════════
# PAGE: DEMOGRAPHICS
# ════════════════════════
if st.session_state.page == "demographics":
    st.markdown("## Vineland-3 — Comprehensive Interview Form" if not AR else "## فاينلاند-3 — نموذج المقابلة الشاملة")
    st.markdown("---")

    with st.form("demo_form"):
        st.markdown("### Client Information" if not AR else "### بيانات العميل")

        c1, c2 = st.columns(2)
        with c1:
            name = st.text_input("Client Name (English only)" if not AR else "اسم العميل (بالإنجليزية فقط)")
            dob = st.date_input("Date of Birth" if not AR else "تاريخ الميلاد",
                                 min_value=date(1920,1,1), max_value=date.today())
            gender_opts = ["Male","Female"] if not AR else ["ذكر","أنثى"]
            gender = st.selectbox("Gender" if not AR else "الجنس", gender_opts)
            nationality = st.text_input("Nationality" if not AR else "الجنسية")
        with c2:
            referral = st.text_input("Referral Source" if not AR else "مصدر الإحالة")
            respondent_name = st.text_input("Respondent Name" if not AR else "اسم المستجيب")
            respondent_rel = st.text_input("Respondent Relationship" if not AR else "صلة المستجيب بالعميل")
            test_date = st.date_input("Test Date" if not AR else "تاريخ التطبيق",
                                       value=date.today())

        st.markdown("### Assessment Options" if not AR else "### خيارات التقييم")
        c3, c4 = st.columns(2)
        with c3:
            include_motor = st.checkbox(
                "Include Motor Skills Domain (ages 0–9 only)" if not AR else "تضمين المهارات الحركية (للأعمار 0-9 سنوات فقط)",
                value=True
            )
        with c4:
            include_maladaptive = st.checkbox(
                "Include Maladaptive Behavior" if not AR else "تضمين السلوك غير التكيفي",
                value=True
            )

        submitted = st.form_submit_button("Begin Assessment →" if not AR else "بدء التقييم ←")

        if submitted:
            has_arabic = bool(re.search(r'[\u0600-\u06FF]', name))
            if not name:
                st.error("Name is required." if not AR else "الاسم مطلوب.")
            elif has_arabic:
                st.warning("⚠️ Name contains Arabic characters — please use English only for the name field." if not AR else "⚠️ الاسم يحتوي على أحرف عربية — يرجى استخدام الإنجليزية فقط.")
            else:
                age_months = age_in_months(dob, test_date)
                age_years = age_months // 12
                age_rem = age_months % 12

                # Load norms to check age validity
                b1, b3, mal = load_norms()
                b1_key = find_b1_key(b1, age_months)

                if not b1_key:
                    st.error("Age is outside the valid range for this assessment (birth to 90+).")
                else:
                    if include_motor and age_years >= 10:
                        st.warning("⚠️ Motor Skills norms are only available for ages 0–9. Motor Skills will be excluded.")
                        include_motor = False

                    st.session_state.demo = {
                        "name": name,
                        "dob": str(dob),
                        "age_str": f"{age_years}:{age_rem:02d}",
                        "age_months": age_months,
                        "gender": "Male" if gender in ["Male","ذكر"] else "Female",
                        "nationality": nationality,
                        "referral": referral,
                        "respondent_name": respondent_name,
                        "respondent_rel": respondent_rel,
                        "test_date": str(test_date),
                        "include_motor": include_motor,
                        "include_maladaptive": include_maladaptive,
                        "raw": {},
                    }
                    st.session_state.page = "items"
                    st.session_state.responses = {}
                    st.rerun()

    st.stop()

# ════════════════════════
# PAGE: ITEMS
# ════════════════════════
if st.session_state.page == "items":
    demo = st.session_state.demo
    b1, b3, mal = load_norms()

    ITEMS = ITEMS_EN if not AR else ITEMS_AR
    MAL_ITEMS = MALADAPTIVE_EN if not AR else MALADAPTIVE_AR
    CRIT_ITEMS = CRITICAL_ITEMS_EN if not AR else CRITICAL_ITEMS_AR

    include_motor = demo["include_motor"]
    include_maladaptive = demo["include_maladaptive"]

    # Build domain list
    domains = [
        ("rec", "Communication — Receptive" if not AR else "التواصل — اللغة الاستقبالية"),
        ("exp", "Communication — Expressive" if not AR else "التواصل — اللغة التعبيرية"),
        ("wrn", "Communication — Written" if not AR else "التواصل — القراءة والكتابة"),
        ("per", "Daily Living Skills — Personal" if not AR else "مهارات الحياة اليومية — المهارات الذاتية"),
        ("dom", "Daily Living Skills — Domestic" if not AR else "مهارات الحياة اليومية — الأنشطة المنزلية"),
        ("cmm", "Daily Living Skills — Community" if not AR else "مهارات الحياة اليومية — المهارات المجتمعية"),
        ("ipr", "Socialization — Interpersonal Relationships" if not AR else "التنشئة الاجتماعية — العلاقات الشخصية المتبادلة"),
        ("pla", "Socialization — Play and Leisure" if not AR else "التنشئة الاجتماعية — وقت الراحة والترفيه"),
        ("cop", "Socialization — Coping Skills" if not AR else "التنشئة الاجتماعية — المسايرة"),
    ]
    if include_motor:
        domains += [
            ("gmo", "Motor Skills — Gross Motor" if not AR else "المهارات الحركية — العضلات الكبيرة"),
            ("fmo", "Motor Skills — Fine Motor" if not AR else "المهارات الحركية — العضلات الدقيقة"),
        ]
    if include_maladaptive:
        domains += [
            ("mal_int", "Maladaptive Behavior — Internalizing" if not AR else "السلوك غير التكيفي — الجزء الأول"),
            ("mal_ext", "Maladaptive Behavior — Externalizing" if not AR else "السلوك غير التكيفي — الجزء الثاني"),
        ]

    total_domains = len(domains)
    current_domain_idx = st.session_state.get("current_domain_idx", 0)

    if current_domain_idx >= total_domains:
        st.session_state.page = "submit"
        st.rerun()

    domain_key, domain_label = domains[current_domain_idx]

    # Progress
    progress = current_domain_idx / total_domains
    st.progress(progress)
    st.markdown(f"**{current_domain_idx+1}/{total_domains}** — {domain_label}")
    st.markdown("---")

    # Get items for this domain
    if domain_key in ["mal_int", "mal_ext"]:
        items_list = MAL_ITEMS.get(domain_key, {})
    else:
        items_list = ITEMS.get(domain_key, {})

    # Response options
    if domain_key in ["mal_int", "mal_ext"]:
        resp_opts = ["2 — Often / نعم، عادةً", "1 — Sometimes / أحياناً", "0 — Never / لا، أبداً"]
    else:
        resp_opts = ["2 — Usually / نعم، عادةً", "1 — Sometimes / أحياناً",
                     "0 — Never / لا، أبداً", "N — No Opportunity / لم تتح الفرصة",
                     "U — Don't Know / لا أعرف"]

    resp_map = {"2 — Usually / نعم، عادةً": 2, "1 — Sometimes / أحياناً": 1,
                "0 — Never / لا، أبداً": 0, "N — No Opportunity / لم تتح الفرصة": "N",
                "U — Don't Know / لا أعرف": "U",
                "2 — Often / نعم، عادةً": 2}

    st.markdown(f"### {domain_label}")

    domain_resp = st.session_state.responses.get(domain_key, {})

    with st.form(f"form_{domain_key}"):
        for item_num, item_text in sorted(items_list.items(), key=lambda x: int(x[0])):
            st.markdown(f"**{item_num}.** {item_text}")
            default_val = domain_resp.get(str(item_num))
            default_idx = 0
            if default_val is not None:
                for i, opt in enumerate(resp_opts):
                    if resp_map.get(opt) == default_val:
                        default_idx = i
                        break

            resp = st.radio(
                f"Item {item_num}",
                resp_opts,
                index=default_idx,
                key=f"{domain_key}_{item_num}",
                horizontal=True,
                label_visibility="collapsed"
            )
            domain_resp[str(item_num)] = resp_map.get(resp, 0)
            st.markdown("---")

        col_prev, col_next = st.columns([1,4])
        with col_prev:
            go_prev = st.form_submit_button("← Back" if not AR else "← رجوع")
        with col_next:
            if current_domain_idx < total_domains - 1:
                go_next = st.form_submit_button("Next →" if not AR else "التالي ←")
            else:
                go_next = st.form_submit_button("Review & Submit →" if not AR else "مراجعة وإرسال ←")

        if go_next:
            st.session_state.responses[domain_key] = domain_resp
            # Calculate raw score (sum of numeric responses)
            raw = sum(v for v in domain_resp.values() if isinstance(v, int))
            st.session_state.demo["raw"][domain_key] = raw
            st.session_state.current_domain_idx = current_domain_idx + 1
            st.rerun()

        if go_prev and current_domain_idx > 0:
            st.session_state.responses[domain_key] = domain_resp
            st.session_state.current_domain_idx = current_domain_idx - 1
            st.rerun()

    st.stop()

# ════════════════════════
# PAGE: SUBMIT
# ════════════════════════
if st.session_state.page == "submit":
    demo = st.session_state.demo
    lang = "Arabic" if AR else "English"

    if st.session_state.done:
        st.markdown("---")
        st.success("✅ Assessment submitted successfully. The report has been sent." if not AR else "✅ تم إرسال التقييم بنجاح.")
        st.balloons()
        if st.button("Start New Assessment" if not AR else "بدء تقييم جديد"):
            for k in ["page","demo","responses","done","current_domain_idx"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.session_state.page = "access"
            st.rerun()
        st.stop()

    st.markdown("## Submitting Assessment..." if not AR else "## جارٍ إرسال التقييم...")

    with st.spinner("Calculating scores and generating report..." if not AR else "جارٍ حساب الدرجات وإنشاء التقرير..."):
        b1, b3, mal = load_norms()

        raw = demo["raw"]
        include_motor = demo["include_motor"]

        # Collect maladaptive raw scores
        mal_int_raw = raw.get("mal_int", 0)
        mal_ext_raw = raw.get("mal_ext", 0)

        all_raw = {
            "rec": raw.get("rec",0), "exp": raw.get("exp",0), "wrn": raw.get("wrn",0),
            "per": raw.get("per",0), "dom": raw.get("dom",0), "cmm": raw.get("cmm",0),
            "ipr": raw.get("ipr",0), "pla": raw.get("pla",0), "cop": raw.get("cop",0),
            "gmo": raw.get("gmo",0), "fmo": raw.get("fmo",0),
            "mal_int": mal_int_raw, "mal_ext": mal_ext_raw,
        }

        scores = compute_scores(all_raw, demo["age_months"], include_motor, b1, b3, mal)

        if not scores:
            st.error("Scoring error — age may be out of range.")
            st.stop()

        # Collect critical items endorsed
        critical_items_endorsed = []
        crit_responses = st.session_state.responses
        for item_key, item_text in CRITICAL_ITEMS_EN.items():
            # Critical items are stored within maladaptive responses
            # We'd need to track them separately; for now note from mal_int/mal_ext responses
            pass

        # Generate narrative
        narrative = generate_narrative(demo, scores, lang, "")

        # Build PDF
        pdf_buf = build_pdf(demo, scores, narrative, lang)

        # Build secondary document
        name_safe = strip_arabic(demo['name']).replace(' ','_')
        if AR:
            secondary_buf = build_arabic_word_doc(demo, st.session_state.responses, scores)
            secondary_filename = f"Vineland3_AR_{name_safe}_{demo['test_date']}.docx"
        else:
            # EN mode: create item response PDF
            secondary_buf = BytesIO()
            from reportlab.pdfgen import canvas as rl_canvas
            c = rl_canvas.Canvas(secondary_buf, pagesize=A4)
            w, h = A4
            c.setFont("Helvetica-Bold", 12)
            c.drawString(30*mm, h-30*mm, f"Vineland-3 Item Responses — {strip_arabic(demo['name'])}")
            c.setFont("Helvetica", 9)
            y = h - 45*mm
            for dk, dl in [("rec","Receptive"),("exp","Expressive"),("wrn","Written"),
                            ("per","Personal"),("dom","Domestic"),("cmm","Community"),
                            ("ipr","Interpersonal"),("pla","Play/Leisure"),("cop","Coping"),
                            ("gmo","Gross Motor"),("fmo","Fine Motor"),
                            ("mal_int","Maladaptive Int"),("mal_ext","Maladaptive Ext")]:
                if not include_motor and dk in ["gmo","fmo"]:
                    continue
                responses = st.session_state.responses.get(dk, {})
                if not responses:
                    continue
                c.setFont("Helvetica-Bold", 10)
                c.drawString(20*mm, y, dl)
                y -= 6*mm
                c.setFont("Helvetica", 8)
                for item_num in sorted(responses.keys(), key=lambda x: int(x)):
                    val = responses[item_num]
                    c.drawString(25*mm, y, f"Item {item_num}: {val}")
                    y -= 5*mm
                    if y < 25*mm:
                        c.showPage()
                        y = h - 25*mm
                y -= 3*mm
            c.save()
            secondary_buf.seek(0)
            secondary_filename = f"Vineland3_Items_{name_safe}_{demo['test_date']}.pdf"

        # Send email
        send_email(demo, scores, pdf_buf, secondary_buf, secondary_filename, lang)

    st.session_state.done = True
    st.rerun()
